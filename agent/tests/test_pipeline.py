from __future__ import annotations

import io
import struct
import zlib

import fitz
import pytest
from docx import Document

import rag.pipeline as pipeline
from core.config import settings
from rag.pipeline import ExtractedAsset


def _png_pixel(red: int, green: int, blue: int) -> bytes:
    def chunk(name: bytes, payload: bytes) -> bytes:
        checksum = zlib.crc32(name + payload) & 0xFFFFFFFF
        return struct.pack(">I", len(payload)) + name + payload + struct.pack(">I", checksum)

    signature = b"\x89PNG\r\n\x1a\n"
    header = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    pixels = zlib.compress(bytes([0, red, green, blue]))
    return signature + chunk(b"IHDR", header) + chunk(b"IDAT", pixels) + chunk(b"IEND", b"")


def test_cleaning_and_cloud_redaction_are_separate() -> None:
    cleaned, stats = pipeline.clean_markdown(
        "Confluence Page ID: 123\n配置 Bearer abcdefghijklmnop\n您的浏览器不支持 video 标签"
    )
    assert "Page ID" not in cleaned
    assert "浏览器不支持" not in cleaned
    assert "Bearer abcdefghijklmnop" in cleaned
    assert stats["confluence_page_id"] == 1
    assert stats["browser_video"] == 1
    assert stats["confluence_page_id_samples"]
    assert "Bearer [REDACTED]" in pipeline.redact_for_cloud(cleaned)


@pytest.mark.parametrize(
    "secret",
    [
        "API_KEY=supersecretvalue123456",
        "api_key: supersecretvalue123456",
        "AWS_ACCESS_KEY_ID=AKIAIOSFODNN7EXAMPLE",
        "AWS_SECRET_ACCESS_KEY=abcdEFGH1234/secret+value",
        "password='correct-horse-battery-staple'",
    ],
)
def test_cloud_redaction_covers_common_secret_assignments(secret: str) -> None:
    redacted = pipeline.redact_for_cloud(secret)
    assert "supersecret" not in redacted
    assert "AKIAIOSFODNN7EXAMPLE" not in redacted
    assert "correct-horse" not in redacted
    assert "[REDACTED" in redacted


def test_image_ocr_fails_closed_without_raw_image_opt_in(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(settings, "vision_api_key", "configured")
    monkeypatch.setattr(settings, "vision_base_url", "https://vision.example/v1")
    monkeypatch.setattr(settings, "vision_allow_raw_images", False)

    def unexpected_post(*_args: object, **_kwargs: object) -> object:
        raise AssertionError("raw image must not be sent to a remote endpoint by default")

    monkeypatch.setattr(pipeline.httpx, "post", unexpected_post)
    assert pipeline.analyze_image(ExtractedAsset(b"secret screenshot", "image/png")) == (
        "",
        "文档图片",
    )


def test_docx_preserves_paragraph_and_table_order() -> None:
    document = Document()
    document.add_paragraph("表格之前")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "参数"
    table.rows[0].cells[1].text = "值"
    table.rows[1].cells[0].text = "timeout"
    table.rows[1].cells[1].text = "30"
    document.add_paragraph("表格之后")
    output = io.BytesIO()
    document.save(output)
    markdown, _assets = pipeline._extract_docx(output.getvalue())
    assert markdown.index("表格之前") < markdown.index("| 参数 | 值 |") < markdown.index("表格之后")


def test_docx_images_keep_document_order_and_heading_context(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    red, blue = _png_pixel(255, 0, 0), _png_pixel(0, 0, 255)
    document = Document()
    document.add_heading("安装", level=1)
    document.add_paragraph("安装步骤").add_run().add_picture(io.BytesIO(red))
    document.add_heading("配置", level=1)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run().add_picture(io.BytesIO(blue))
    document.add_heading("排错", level=1)
    document.add_picture(io.BytesIO(red))
    output = io.BytesIO()
    document.save(output)

    markdown, assets = pipeline._extract_docx(output.getvalue())

    assert "# 安装" in markdown and "# 配置" in markdown and "# 排错" in markdown
    assert [asset.data for asset in assets] == [red, blue, red]
    assert [asset.heading_path for asset in assets] == ["安装", "配置", "排错"]

    monkeypatch.setattr(pipeline, "extract_document", lambda *_args: (markdown, assets))
    monkeypatch.setattr(pipeline, "write_derived", lambda *_args: None)
    monkeypatch.setattr(pipeline, "analyze_image", lambda _asset: ("OCR", "说明"))
    result = pipeline.process_document(1, 1, "", "docx", "source")
    image_chunks = [chunk for chunk in result.chunks if chunk.kind == "image"]

    assert len(result.assets) == 2
    assert result.assets[0].heading_path == "安装 | 排错"
    assert [chunk.heading_path for chunk in image_chunks] == ["安装 | 排错", "配置"]


def test_pdf_keeps_page_numbers() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "MQTT config")
    blob = document.tobytes()
    document.close()
    markdown, _assets = pipeline._extract_pdf(blob)
    assert "第 1 页" in markdown
    assert "MQTT config" in markdown


def test_long_chunks_respect_budget(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(settings, "max_chunk_tokens", 100)
    monkeypatch.setattr(settings, "chunk_overlap_tokens", 10)
    monkeypatch.setattr(settings, "short_document_chars", 10)
    chunks = pipeline.chunk_body("\n\n".join(["配置参数 " * 20 for _ in range(10)]))
    assert len(chunks) > 1
    assert all(chunk.token_count <= 120 for chunk in chunks)


def test_assets_are_sha_deduplicated(monkeypatch: pytest.MonkeyPatch) -> None:
    duplicate = ExtractedAsset(data=b"same-image", mime_type="image/png")
    monkeypatch.setattr(
        pipeline, "extract_document", lambda *_args: ("正文", [duplicate, duplicate])
    )
    monkeypatch.setattr(pipeline, "write_derived", lambda *_args: None)
    monkeypatch.setattr(pipeline, "analyze_image", lambda _asset: ("OCR", "说明"))
    result = pipeline.process_document(1, 1, "", "docx", "source")
    assert len(result.assets) == 1
    assert result.assets[0].storage_key.startswith("assets/")
    assert len([chunk for chunk in result.chunks if chunk.kind == "image"]) == 1
