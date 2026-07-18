"""RAG v2 文档提取、清洗、脱敏、语义增强、分块与图片处理流水线。"""

from __future__ import annotations

import base64
import hashlib
import io
import json
import logging
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path

import fitz
import httpx
import yaml
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.xmlchemy import BaseOxmlElement

from db import settings
from storage import read_source, write_derived

StageCallback = Callable[[str, dict[str, int | str]], None]
logger = logging.getLogger("agent.pipeline")


@dataclass
class ExtractedAsset:
    data: bytes
    mime_type: str
    page_number: int | None = None
    sha256: str = ""
    storage_key: str = ""
    ocr_text: str = ""
    caption: str = ""
    width: int | None = None
    height: int | None = None
    heading_path: str = ""


@dataclass
class ChunkRecord:
    kind: str
    chunk_idx: int
    content: str
    token_count: int
    lexical_text: str
    heading_path: str = ""
    page_number: int | None = None
    asset_index: int | None = None


@dataclass
class PipelineResult:
    markdown: str
    normalized_storage_key: str
    summary: str
    keywords: list[str]
    questions: list[str]
    cleaning_stats: dict[str, int | list[str]]
    chunks: list[ChunkRecord] = field(default_factory=list)
    assets: list[ExtractedAsset] = field(default_factory=list)


def redact_for_cloud(text: str) -> str:
    """仅遮蔽发送到云模型的敏感片段，本地规范正文保持原内容。"""
    patterns = [
        (r"(?i)\bBearer\s+[A-Za-z0-9._~+/=-]{12,}", "Bearer [REDACTED]"),
        (
            r"(?i)\b((?:(?:aws[_-]?)?(?:access[_-]?key(?:[_-]?id)?|"
            r"secret[_-]?access[_-]?key|session[_-]?token)|api[_-]?key|secret(?:[_-]?key)?|"
            r"auth[_-]?token|refresh[_-]?token|password|passwd|pwd)\s*[:=]\s*)"
            r"[\"']?[A-Za-z0-9_./+=:@-]{8,}[\"']?",
            r"\1[REDACTED]",
        ),
        (r"\bAKIA[0-9A-Z]{16}\b", "[REDACTED_AWS_ACCESS_KEY]"),
        (r"(?i)\b(?:sk|api)[-_][A-Za-z0-9_-]{12,}\b", "[REDACTED_KEY]"),
        (r"\b[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}\b", "[REDACTED_EMAIL]"),
        (r"(?<!\d)1[3-9]\d{9}(?!\d)", "[REDACTED_PHONE]"),
    ]
    redacted = text
    for expression, replacement in patterns:
        redacted = re.sub(expression, replacement, redacted)
    return redacted


def _rules() -> tuple[str, list[tuple[str, re.Pattern[str]]]]:
    path = Path(__file__).with_name("cleaning_rules.yml")
    payload = yaml.safe_load(path.read_text(encoding="utf-8"))
    compiled = [
        (item["name"], re.compile(item["expression"])) for item in payload.get("patterns", [])
    ]
    return str(payload.get("version", settings.cleaning_rules_version)), compiled


def clean_markdown(text: str) -> tuple[str, dict[str, int | list[str]]]:
    """应用可审计规则并收敛空白，不吞掉正文结构。"""
    _version, rules = _rules()
    cleaned = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    stats: dict[str, int | list[str]] = {}
    for name, pattern in rules:
        samples = [match.group(0).strip()[:120] for match in list(pattern.finditer(cleaned))[:3]]
        cleaned, count = pattern.subn("", cleaned)
        stats[name] = count
        if samples:
            stats[f"{name}_samples"] = samples
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned, stats


def estimate_tokens(text: str) -> int:
    """无额外 tokenizer 依赖的保守估算；中文按字、英文按词计数。"""
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    latin = len(re.findall(r"[A-Za-z0-9_./:-]+", text))
    return cjk + latin


def _docx_assets_in_element(
    element: BaseOxmlElement, document: DocxDocument, heading_path: str
) -> list[ExtractedAsset]:
    """按 XML 出现顺序提取当前段落或表格中的内嵌图片。"""
    assets: list[ExtractedAsset] = []
    for blip in element.xpath(".//a:blip"):
        relation_id = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if not relation_id or relation_id not in document.part.rels:
            continue
        relation = document.part.rels[relation_id]
        if "image" not in relation.reltype:
            continue
        part = relation.target_part
        assets.append(
            ExtractedAsset(
                data=part.blob,
                mime_type=part.content_type,
                heading_path=heading_path,
            )
        )
    return assets


def _extract_docx(blob: bytes) -> tuple[str, list[ExtractedAsset]]:
    document = Document(io.BytesIO(blob))
    blocks: list[str] = []
    assets: list[ExtractedAsset] = []
    heading_stack: list[str] = []
    # 遍历 body 子节点，避免把所有表格错误地移到正文末尾。
    paragraphs = iter(document.paragraphs)
    tables = iter(document.tables)
    for element in document.element.body.iterchildren():
        if element.tag.endswith("}p"):
            paragraph = next(paragraphs)
        elif element.tag.endswith("}tbl"):
            table = next(tables)
            rows = [
                [cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows
            ]
            if rows:
                blocks.append("| " + " | ".join(rows[0]) + " |")
                blocks.append("| " + " | ".join("---" for _ in rows[0]) + " |")
                blocks.extend("| " + " | ".join(row) + " |" for row in rows[1:])
            assets.extend(_docx_assets_in_element(element, document, " > ".join(heading_stack)))
            continue
        else:
            continue
        text = paragraph.text.strip()
        style = paragraph.style.name.lower() if paragraph.style else ""
        if text and style.startswith("heading"):
            level = re.findall(r"\d+", style)
            heading_level = int(level[0] if level else 2)
            heading_stack = heading_stack[: max(0, heading_level - 1)]
            heading_stack.append(text)
            blocks.append(f"{'#' * heading_level} {text}")
        elif text:
            blocks.append(text)
        assets.extend(_docx_assets_in_element(element, document, " > ".join(heading_stack)))
    return "\n\n".join(blocks), assets


def _extract_pdf(blob: bytes) -> tuple[str, list[ExtractedAsset]]:
    document = fitz.open(stream=blob, filetype="pdf")
    pages: list[str] = []
    assets: list[ExtractedAsset] = []
    for page_index, page in enumerate(document):
        text = page.get_text("text").strip()
        if estimate_tokens(text) < 20:
            # 扫描页或低文本密度页渲染为整页图片交给 OCR/VL；正文提取失败不阻断入库。
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            assets.append(
                ExtractedAsset(
                    data=pixmap.tobytes("png"),
                    mime_type="image/png",
                    page_number=page_index + 1,
                    heading_path=f"第 {page_index + 1} 页",
                )
            )
        try:
            for table in page.find_tables().tables:
                rows = table.extract()
                if rows:
                    text += "\n\n| " + " | ".join(str(cell or "") for cell in rows[0]) + " |"
                    text += "\n| " + " | ".join("---" for _ in rows[0]) + " |"
                    text += "".join(
                        "\n| " + " | ".join(str(cell or "") for cell in row) + " |"
                        for row in rows[1:]
                    )
        except Exception:
            pass
        pages.append(
            f"## 第 {page_index + 1} 页\n\n{text}" if text else f"## 第 {page_index + 1} 页"
        )
        for image in page.get_images(full=True):
            extracted = document.extract_image(image[0])
            assets.append(
                ExtractedAsset(
                    data=extracted["image"],
                    mime_type=f"image/{extracted['ext']}",
                    page_number=page_index + 1,
                    heading_path=f"第 {page_index + 1} 页",
                )
            )
    document.close()
    return "\n\n".join(pages), assets


def extract_document(
    file_type: str, content: str, storage_key: str
) -> tuple[str, list[ExtractedAsset]]:
    """按受支持格式提取规范文本与图片；文本请求保持旧接口兼容。"""
    kind = (file_type or "txt").lower().lstrip(".")
    if kind in {"txt", "md", "markdown"} and content:
        return content, []
    blob = read_source(storage_key) if storage_key else content.encode()
    if kind in {"txt", "md", "markdown"}:
        return blob.decode("utf-8", errors="replace"), []
    if kind == "docx":
        return _extract_docx(blob)
    if kind == "pdf":
        return _extract_pdf(blob)
    raise ValueError(f"unsupported file type: {kind}")


def _fallback_metadata(text: str) -> tuple[str, list[str], list[str]]:
    first = re.split(r"[。！？!?\n]", text.strip())[0][:180] or "资料内容"
    terms = re.findall(r"[A-Za-z][A-Za-z0-9_.-]{2,}|[\u4e00-\u9fff]{2,8}", text)
    keywords = list(dict.fromkeys(terms))[:8]
    if len(keywords) < 5:
        keywords.extend([f"主题{i}" for i in range(1, 6 - len(keywords))])
    questions = [f"{term}是什么？" for term in keywords[:3]]
    return first, keywords[:8], questions


def enrich_metadata(text: str) -> tuple[str, list[str], list[str]]:
    """生成摘要、关键词和候选问题；云调用失败时确定性降级。"""
    if not settings.llm_api_key or not settings.llm_base_url:
        return _fallback_metadata(text)
    prompt = (
        "阅读资料并只输出 JSON：summary 为一句话；keywords 为 5 到 8 个技术术语；"
        "questions 为 3 到 5 个该资料能回答的问题。\n\n" + redact_for_cloud(text[:30000])
    )
    try:
        response = httpx.post(
            f"{settings.llm_base_url.rstrip('/')}/chat/completions",
            headers={"Authorization": f"Bearer {settings.llm_api_key}"},
            json={
                "model": settings.llm_model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0,
                "response_format": {"type": "json_object"},
            },
            timeout=settings.parser_embedding_timeout_s,
        )
        response.raise_for_status()
        result = json.loads(response.json()["choices"][0]["message"]["content"])
        summary = str(result["summary"]).strip()[:500]
        keywords = [str(item).strip()[:80] for item in result["keywords"] if str(item).strip()][:8]
        questions = [str(item).strip()[:300] for item in result["questions"] if str(item).strip()][
            :5
        ]
        if not summary or len(keywords) < 5 or len(questions) < 3:
            raise ValueError("incomplete metadata")
        return summary, keywords, questions
    except Exception:
        return _fallback_metadata(text)


def analyze_image(asset: ExtractedAsset) -> tuple[str, str]:
    """调用 Qwen OCR/VL 生成图片 OCR 与说明，失败时保留可展示资产。"""
    api_key = settings.vision_api_key or settings.embedding_api_key
    base_url = settings.vision_base_url or settings.embedding_base_url
    if not api_key or not base_url or not settings.vision_allow_raw_images:
        if api_key and base_url and not settings.vision_allow_raw_images:
            logger.info("vision OCR skipped", extra={"reason": "raw_images_disabled"})
        return "", "文档图片"
    encoded = base64.b64encode(asset.data).decode()
    try:
        response = httpx.post(
            f"{base_url.rstrip('/')}/chat/completions",
            headers={"Authorization": f"Bearer {api_key}"},
            json={
                "model": settings.vision_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:{asset.mime_type};base64,{encoded}"},
                            },
                            {
                                "type": "text",
                                "text": "输出 JSON：ocr_text 为图片文字，caption 为图片对技术文档的简短说明。",
                            },
                        ],
                    }
                ],
                "response_format": {"type": "json_object"},
            },
            timeout=settings.parser_embedding_timeout_s,
        )
        response.raise_for_status()
        result = json.loads(response.json()["choices"][0]["message"]["content"])
        return str(result.get("ocr_text", ""))[:12000], str(result.get("caption", ""))[:1000]
    except Exception:
        return "", "文档图片"


def chunk_body(markdown: str) -> list[ChunkRecord]:
    """短文档整篇入库；长文档按段落递归组成 token 有界块。"""
    total_tokens = estimate_tokens(markdown)
    if len(markdown) <= settings.short_document_chars and total_tokens <= 3500:
        pages = re.findall(r"(?m)^#{1,6}\s+第\s*(\d+)\s*页\s*$", markdown)
        return [
            ChunkRecord(
                "body",
                0,
                markdown,
                total_tokens,
                markdown,
                page_number=int(pages[0]) if pages else None,
            )
        ]
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", markdown) if part.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_tokens = 0
    for paragraph in paragraphs:
        paragraph_tokens = estimate_tokens(paragraph)
        if current and current_tokens + paragraph_tokens > settings.max_chunk_tokens:
            chunks.append("\n\n".join(current))
            overlap: list[str] = []
            overlap_tokens = 0
            for previous in reversed(current):
                previous_tokens = estimate_tokens(previous)
                if overlap_tokens + previous_tokens > settings.chunk_overlap_tokens:
                    break
                overlap.insert(0, previous)
                overlap_tokens += previous_tokens
            current = overlap
            current_tokens = overlap_tokens
        if paragraph_tokens > settings.max_chunk_tokens:
            if current:
                chunks.append("\n\n".join(current))
                current, current_tokens = [], 0
            # 字符上限不大于 token 上限，对中英文都是保守边界；超长无换行表格也不会越界。
            step = settings.max_chunk_tokens
            overlap_chars = min(settings.chunk_overlap_tokens, max(0, step - 1))
            cursor = 0
            while cursor < len(paragraph):
                chunks.append(paragraph[cursor : cursor + step])
                cursor += max(1, step - overlap_chars)
            continue
        current.append(paragraph)
        current_tokens += paragraph_tokens
    if current:
        chunks.append("\n\n".join(current))
    records: list[ChunkRecord] = []
    cursor = 0
    for index, chunk in enumerate(chunks):
        if not chunk.strip():
            continue
        probe = chunk[: min(80, len(chunk))]
        position = markdown.find(probe, max(0, cursor - settings.chunk_overlap_tokens * 2))
        if position < 0:
            position = markdown.find(probe)
        prefix = markdown[: max(0, position)]
        headings = re.findall(r"(?m)^#{1,6}\s+(.+?)\s*$", prefix)
        pages = re.findall(r"(?m)^#{1,6}\s+第\s*(\d+)\s*页\s*$", prefix)
        records.append(
            ChunkRecord(
                "body",
                index,
                chunk,
                estimate_tokens(chunk),
                chunk,
                heading_path=headings[-1] if headings else "",
                page_number=int(pages[-1]) if pages else None,
            )
        )
        cursor = max(cursor, position + len(chunk))
    return records


def process_document(
    material_id: int,
    generation: int,
    content: str,
    file_type: str,
    storage_key: str,
    stage_callback: StageCallback | None = None,
) -> PipelineResult:
    """执行完整解析流水线并产出可原子持久化的 RAG v2 结果。"""
    notify = stage_callback or (lambda _stage, _progress: None)
    notify("extract", {})
    raw_text, assets = extract_document(file_type, content, storage_key)
    notify("clean", {"raw_chars": len(raw_text)})
    markdown, cleaning_stats = clean_markdown(raw_text)
    notify("enrich", {"clean_chars": len(markdown)})
    summary, keywords, questions = enrich_metadata(markdown)
    notify("assets", {"asset_count": len(assets)})
    unique_assets: list[ExtractedAsset] = []
    seen_assets: dict[str, ExtractedAsset] = {}
    for asset in assets:
        asset.sha256 = hashlib.sha256(asset.data).hexdigest()
        existing = seen_assets.get(asset.sha256)
        if existing is not None:
            paths = [part for part in existing.heading_path.split(" | ") if part]
            if asset.heading_path and asset.heading_path not in paths:
                existing.heading_path = " | ".join([*paths, asset.heading_path])
            continue
        seen_assets[asset.sha256] = asset
        extension = asset.mime_type.split("/")[-1].replace("jpeg", "jpg")
        asset.storage_key = f"assets/{asset.sha256}.{extension}"
        try:
            pixmap = fitz.Pixmap(asset.data)
            asset.width, asset.height = pixmap.width, pixmap.height
        except Exception:
            pass
        asset.ocr_text, asset.caption = analyze_image(asset)
        try:
            write_derived(asset.storage_key, asset.data, asset.mime_type)
        except Exception as exc:
            logger.warning("asset storage degraded", extra={"error_type": type(exc).__name__})
            asset.storage_key = ""
        unique_assets.append(asset)
    assets = unique_assets
    normalized_key = f"{material_id}/{generation}/normalized.md"
    try:
        write_derived(normalized_key, markdown.encode(), "text/markdown; charset=utf-8")
    except Exception as exc:
        logger.warning(
            "normalized document storage degraded", extra={"error_type": type(exc).__name__}
        )
        normalized_key = ""
    notify("chunk", {})
    chunks = chunk_body(markdown)
    chunks.append(ChunkRecord("summary", 0, summary, estimate_tokens(summary), summary))
    for index, question in enumerate(questions):
        chunks.append(ChunkRecord("question", index, question, estimate_tokens(question), question))
    for index, asset in enumerate(assets):
        signal = "\n".join(part for part in (asset.caption, asset.ocr_text) if part).strip()
        if signal:
            chunks.append(
                ChunkRecord(
                    "image",
                    index,
                    signal,
                    estimate_tokens(signal),
                    signal,
                    page_number=asset.page_number,
                    asset_index=index,
                    heading_path=asset.heading_path,
                )
            )
    return PipelineResult(
        markdown=markdown,
        normalized_storage_key=normalized_key,
        summary=summary,
        keywords=keywords,
        questions=questions,
        cleaning_stats=cleaning_stats,
        chunks=chunks,
        assets=assets,
    )
